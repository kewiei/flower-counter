
import re
import docx as docx
import xlwt


class FlowerOrder:

    def __init__(self):
        self.id = -1
        self.name = ''
        self.items = []
        self.exceptions = []

    def addid(self, id):
        self.id = id

    def addname(self, name):
        self.name = name

    def additem(self, item):
        self.items.append(item)

    def addexception(self, exception):
        self.exceptions.append(exception)


def isnumber(str):
    for j in range(len(str)):
        if str[j] not in ('0', '1', '2', '3', '4', '5', '6', '7', '8', '9'):
            return False
    if str[0] == '0':
        return False
    else:
        return True


def getflowername(str):
    for j in range(len(str)):
        if str[j] in ('0', '1', '2', '3', '4', '5', '6', '7', '8', '9'):
            return str[:j]


def getflowernum(str):
    for j in range(len(str)):
        if str[j] in ('0', '1', '2', '3', '4', '5', '6', '7', '8', '9'):
            return str[j:]


def process(sourcepath, outputpath):
    # 获取文档对象
    file = docx.Document(sourcepath)
    print("段落数:" + str(len(file.paragraphs)))

    book = xlwt.Workbook('utf-8', 0)
    sheet = book.add_sheet('汇总', True)

    orders = []

    # 输出段落编号及段落内容
    for i in range(len(file.paragraphs)):
        print("第" + str(i) + "段的内容是：" + file.paragraphs[i].text)
        temppara = file.paragraphs[i].text
        splited = re.split('[ ,.;:，：、]', temppara)

        # 去除空项和空格项
        i = 0
        while i < len(splited):
            splited[i] = splited[i].strip().strip('\xa0')
            if splited[i] == ' ' or splited[i] == '' or splited[i] == '\xa0':
                splited.pop(i)
                i -= 1
            i += 1
        # 为了去除空段
        if len(splited) == 0:
            print(">>>>>发现空行")
            continue

        # 为了去除首项中数字和微信名字的连接
        if not splited[0].endswith(('0', '1', '2', '3', '4', '5', '6', '7', '8', '9')):
            stuck = splited[0]
            i = len(stuck) - 1
            while i >= 0:
                if stuck[i] in ('0', '1', '2', '3', '4', '5', '6', '7', '8', '9'):
                    splited[0] = stuck[:i + 1]
                    splited.insert(1, stuck[i + 1:])
                    break
                i -= 1
            print(">>>>>数字和微信名字的连接已去除")

        # 为了让单独的数字与之前结合
        i = 3
        while i < len(splited):
            if isnumber(splited[i]):
                splited[i - 1] = splited[i - 1] + splited[i]
                splited.pop(i)
                print(">>>>>单独的数字与之前结合完成")
                i -= 1
            i += 1

        # 为了让每一项开头的数字与之前结合
        i = 3
        while i < len(splited):
            if splited[i][0] in ('1', '2', '3', '4', '5', '6', '7', '8', '9'):
                j = 0
                while j < len(splited[i]):
                    if splited[i][j] not in ('1', '2', '3', '4', '5', '6', '7', '8', '9'):
                        temp = splited[i]
                        splited[i - 1] += temp[:j]
                        splited[i] = temp[j:]
                        break
                    j += 1
                print(">>>>>每一项开头的数字与之前结合完成")
            i += 1

        # 去除"扎"
        i = 2
        while i < len(splited):
            if splited[i].find("扎") != -1:
                splited[i] = splited[i].strip('扎')
                print(">>>>>已删除 扎")
            i += 1

        # 将连到一起的数据分开
        i = 2
        while i < len(splited):
            j = 0
            while j < len(splited[i]):
                if splited[i][j] in ('0', '1', '2', '3', '4', '5', '6', '7', '8', '9') \
                        and j < len(splited[i]) - 1 \
                        and splited[i][j + 1] not in ('0', '1', '2', '3', '4', '5', '6', '7', '8', '9'):
                    splited.insert(i + 1, splited[i][j + 1:])
                    splited[i] = splited[i][:j + 1]
                    print(">>>>>已将连到一起的数据分开")
                j += 1
            i += 1

        # 数据分发
        tempOrder = FlowerOrder()
        for j in range(len(splited)):
            if j == 0:
                print("编号:", splited[j])
                tempOrder.addid(splited[j])
            elif j == 1:
                print("名字:", splited[j])
                tempOrder.addname(splited[j])
            elif splited[j].endswith(('0', '1', '2', '3', '4', '5', '6', '7', '8', '9')):
                print(splited[j])
                tempOrder.additem(splited[j])
            else:
                print("异常:", splited[j])
                tempOrder.addexception(splited[j])

        orders.append(tempOrder)

    flowermentioned = []
    for order in orders:
        for item in order.items:
            if not getflowername(item) in flowermentioned:
                flowermentioned.append(getflowername(item))

    flowermentioned.sort(reverse=False)

    for i in range(len(flowermentioned)):
        sheet.write(0, i + 1, flowermentioned[i])

    for i in range(len(orders)):
        # 写名字
        sheet.write(i + 1, 0, orders[i].name)
        for item in orders[i].items:
            sheet.write(i + 1, flowermentioned.index(getflowername(item)) + 1, int(getflowernum(item)))

    # 最后，将以上操作保存到指定的Excel文件中
    book.save(outputpath)
    # 书写例外
    outputexception = ''
    for i in range(len(orders)):
        if len(orders[i].exceptions) > 0:
            outputexception += "序号" + orders[i].id + '\n'
            for exception in orders[i].exceptions:
                outputexception += exception + '\n'
    print("输出的例外\n")
    print(outputexception)
    return outputexception

if __name__ == '__main__':
    process("e:\\flower.docx", r'e:\outcome.xls')
